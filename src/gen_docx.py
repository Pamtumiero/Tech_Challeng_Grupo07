# -*- coding: utf-8 -*-
"""Gera o relatorio .docx (ABNT + identidade FIAP) da analise Geografia e Categorias."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

ROOT = Path(__file__).resolve().parent.parent
FIG = ROOT / "figs"
OUT = ROOT / "apresentacao" / "Relatorio_Tech_Challenge_Olist_ABNT_FIAP.docx"

MAGENTA = RGBColor(0xED, 0x14, 0x5B)   # FIAP Magenta
PRETO   = RGBColor(0x1A, 0x1A, 0x1A)
FONTE   = "Arial"

doc = Document()

# ---------- estilo Normal (ABNT: Arial 12, 1.5, justificado) ----------
normal = doc.styles["Normal"]
normal.font.name = FONTE; normal.font.size = Pt(12); normal.font.color.rgb = PRETO
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.space_after = Pt(0)

# ---------- pagina A4 e margens ABNT (sup/esq 3cm, inf/dir 2cm) ----------
sec = doc.sections[0]
sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
sec.top_margin = Cm(3); sec.left_margin = Cm(3); sec.bottom_margin = Cm(2); sec.right_margin = Cm(2)

_fig = {"n": 0}

def _set_font(run, size=12, bold=False, italic=False, color=PRETO, font=FONTE):
    run.font.name = font; run.font.size = Pt(size); run.font.bold = bold
    run.font.italic = italic; run.font.color.rgb = color

def centro(text, size=12, bold=False, color=PRETO, before=0, after=0, italic=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for i, ln in enumerate(text.split("\n")):
        if i: p.add_run().add_break()
        _set_font(p.add_run(ln), size, bold, italic, color)
    return p

def h(text, size=13, before=14, after=8):
    """Titulo de secao numerado, magenta FIAP, sem recuo."""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _set_font(p.add_run(text), size, bold=True, color=MAGENTA)
    return p

def para(text, indent=True):
    p = doc.add_paragraph(text)
    if indent: p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.25); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for r in p.runs: _set_font(r, 11)
    return p

def figura(arquivo, titulo, fonte="Elaborado pelos autores (2026), a partir do dataset Olist (Kaggle)."):
    _fig["n"] += 1
    p = FIG / arquivo
    iw, ih = Image.open(p).size
    larg = Cm(15.0); alt = Cm(15.0 * ih / iw)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(8); cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _set_font(cap.add_run(f"Figura {_fig['n']} — {titulo}"), 10, bold=True)
    pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.paragraph_format.space_after = Pt(0)
    pi.add_run().add_picture(str(p), width=larg, height=alt)
    pf2 = doc.add_paragraph(); pf2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf2.paragraph_format.space_after = Pt(10); pf2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _set_font(pf2.add_run(f"Fonte: {fonte}"), 10)

def tabela(headers, rows, larguras_cm):
    _fig  # noop
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htxt); _set_font(run, 10, bold=True, color=RGBColor(255,255,255))
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade(hdr[i], "ED145B")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val)); _set_font(run, 10)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in t.rows:
        for i, w in enumerate(larguras_cm):
            row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def page_break():
    doc.add_page_break()

def add_page_numbers():
    """Numero de pagina no canto superior direito (ABNT)."""
    hdr = sec.header
    p = hdr.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    _set_font(run, 10)

def add_sumario():
    p = doc.add_paragraph()
    run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = r'TOC \o "1-3" \h \z \u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Atualize o sumário (clique com o botão direito › Atualizar campo)."
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(it); run._r.append(f2); run._r.append(t); run._r.append(f3)

# =====================================================================
# CAPA (ABNT)
# =====================================================================
centro("FACULDADE DE INFORMÁTICA E ADMINISTRAÇÃO PAULISTA — FIAP", 12, bold=True, before=6, after=2)
centro("Pós-Graduação (POSTECH) — Data Analytics", 12, after=120)
centro("GEOGRAFIA E CATEGORIAS:\nDESTRAVAR O BRASIL ALÉM DO SUDESTE", 16, bold=True, color=MAGENTA, after=10)
centro("Diagnóstico geográfico-logístico do marketplace Olist e proposta de descentralização da oferta",
       12, italic=True, after=170)
centro("Tech Challenge — Fase 1", 12, bold=True, after=2)
centro("São Paulo\n2026", 12, bold=True)
page_break()

# =====================================================================
# FOLHA DE ROSTO
# =====================================================================
centro("Lucas Aguiar de Moura — RM375551\nPamela Regina Tumiero da Costa — RM374909\n"
       "Guilherme Augusto Justino da Silva — RM374370\nMariana Nogueira Salgado Zeron — RM374774\n"
       "Vitor Martino — RM374436", 12, bold=True, before=6, after=90)
centro("GEOGRAFIA E CATEGORIAS:\nDESTRAVAR O BRASIL ALÉM DO SUDESTE", 14, bold=True, color=MAGENTA, after=40)
nota = doc.add_paragraph()
nota.paragraph_format.left_indent = Cm(8); nota.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
nota.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
_set_font(nota.add_run("Relatório técnico-analítico apresentado como requisito de avaliação do Tech Challenge "
    "Fase 1 do curso de Pós-Graduação em Data Analytics da FIAP, com base no Brazilian E-Commerce Public "
    "Dataset by Olist."), 11)
doc.add_paragraph().paragraph_format.space_after = Pt(120)
centro("São Paulo\n2026", 12, bold=True)
page_break()

# =====================================================================
# RESUMO
# =====================================================================
h("RESUMO", size=13, before=4)
para("Este relatório investiga o problema geográfico do marketplace Olist: por que as regiões fora do Sudeste "
     "compram menos e recebem com mais atraso, e o que fazer a respeito. Sobre 96.478 pedidos efetivamente "
     "entregues (set/2016–ago/2018), demonstra-se que a concentração da oferta no Sudeste (64,6% da receita) é a "
     "causa raiz comum de dois sintomas: (i) o frete elevado — mediana de 31,5% do preço no Nordeste — filtra a "
     "compra de menor valor, e (ii) a perna de transporte de longa distância estoura o prazo prometido. Mostra-se, "
     "por decomposição do lead time e por teste de hipótese (Welch t-test, p≈0), que o gargalo é o transporte — não "
     "o vendedor — e que, no mesmo estado, o Nordeste entrega tão rápido quanto o Sudeste (7 dias). Conclui-se com "
     "uma proposta de descentralização da oferta (fulfillment distribuído nos estados-âncora e recrutamento de "
     "sellers locais), cujo cenário de 30% projeta ~169 atrasos evitados e ~R$ 51 mil de economia de frete.", indent=False)
pal = doc.add_paragraph(); pal.paragraph_format.space_before = Pt(8)
_set_font(pal.add_run("Palavras-chave: "), 12, bold=True)
_set_font(pal.add_run("e-commerce; logística; análise geográfica; satisfação do cliente; Olist."), 12)
page_break()

# =====================================================================
# SUMÁRIO
# =====================================================================
h("SUMÁRIO", size=13, before=4)
add_sumario()
page_break()

# =====================================================================
# 1. INTRODUÇÃO
# =====================================================================
h("1 INTRODUÇÃO")
para("A Olist é um integrador de marketplaces que conecta milhares de pequenos e médios vendedores (sellers) "
     "aos grandes portais de e-commerce do Brasil. Em aproximadamente dois anos, a plataforma movimentou cerca de "
     "R$ 13,2 milhões em mercadoria sobre 96.478 pedidos entregues. Apesar do crescimento expressivo, a operação é "
     "fortemente concentrada na Região Sudeste, o que levanta a questão central deste trabalho: como destravar o "
     "potencial de consumo das demais regiões do país.")
para("Duas perguntas de negócio orientam a análise: (1) por que Nordeste e Centro-Oeste compram menos que o Sul? "
     "e (2) o que pode ser melhorado para cumprir as promessas de prazo de entrega? Demonstra-se que ambas "
     "compartilham a mesma causa raiz — a localização da oferta —, o que conduz a uma recomendação única e mensurável.")

# =====================================================================
# 2. METODOLOGIA
# =====================================================================
h("2 METODOLOGIA")
para("A base utilizada é o Brazilian E-Commerce Public Dataset by Olist (Kaggle), composto por nove tabelas "
     "relacionais (pedidos, itens, pagamentos, avaliações, produtos, clientes, vendedores, geolocalização e "
     "tradução de categorias). O processamento foi conduzido em Python 3.12 (pandas, matplotlib, scipy), de forma "
     "integralmente reprodutível.")
para("Definições e decisões de escopo adotadas:", indent=False)
bullet("Escopo: utilizam-se exclusivamente os pedidos com status delivered (entregues). Os não entregues são "
       "tratados separadamente em adendo e excluídos das métricas, por não possuírem data de entrega, frete "
       "realizado nem avaliação confiáveis.")
bullet("Ticket médio: valor da mercadoria (price), sem o frete — métrica de comportamento de compra.")
bullet("Receita/faturamento: price + frete (GMV recebido pela plataforma).")
bullet("Frete (%): mediana da razão frete/preço calculada por pedido (reflete o pedido típico, robusta a outliers).")
bullet("Macrorregiões definidas pela UF do cliente; análise local×não-local definida pela UF do vendedor.")

# =====================================================================
# 3. PANORAMA
# =====================================================================
h("3 PANORAMA: A CONCENTRAÇÃO NO SUDESTE")
para("A receita é hiperconcentrada no Sudeste (64,6%), e o ticket médio de mercadoria cresce à medida que a região "
     "se afasta do Sudeste — primeiro indício de que algo filtra a compra nas regiões distantes (Figura 1).")
figura("geo_00_panorama.png", "Participação na receita e ticket médio (mercadoria) por região")
para("Do lado das categorias, a receita distribui-se por uma cauda longa saudável — 18 categorias respondem por "
     "80% do faturamento, sem dependência de um único nicho (Figura 2).")
figura("geo_01_categorias.png", "Categorias com maior receita (concentração saudável)")

# =====================================================================
# 4. Q1
# =====================================================================
h("4 POR QUE NORDESTE E CENTRO-OESTE COMPRAM MENOS (Q1)")
para("A resposta não é falta de demanda, e sim falta de oferta local. O Nordeste conta com apenas 56 vendedores "
     "locais e o Centro-Oeste com 79, contra 668 no Sul e 2.287 no Sudeste. Como quase tudo é despachado de longe, "
     "o frete torna-se um pedágio: no pedido típico equivale a ~31% do preço no Nordeste e ~25% no Centro-Oeste, "
     "contra ~20% no Sudeste (Figura 3).")
figura("geo_q1_oferta_frete.png", "Oferta local escassa (esq.) e o frete filtrando a compra (dir.)")
para("O efeito é um filtro de compra: NE e CO exibem o maior ticket de mercadoria do país (R$ 165 e R$ 150) não "
     "por maior poder aquisitivo, mas porque só vale a pena pagar o frete em pedidos caros — as compras pequenas, "
     "que dão volume ao Sul, deixam de acontecer.")

# =====================================================================
# 5. CATEGORIAS x REGIAO
# =====================================================================
h("5 CATEGORIAS POR REGIÃO")
para("As regiões têm preferências distintas, e a leitura confirma a tese do frete. O Nordeste sub-indexa "
     "fortemente em categorias volumosas — cama, mesa e banho (−45%) e utilidades domésticas (−39%) — nas quais o "
     "frete é proibitivo, e sobre-indexa em itens pequenos e de alto valor — saúde e beleza (+41%), relógios e "
     "presentes (+22%) e acessórios automotivos (+28%) —, que “compensam” o frete (Figura 4).")
figura("geo_02_categoria_overindex.png", "Sobre/sub-indexação de categorias por região vs. média nacional")

# =====================================================================
# 6. Q2
# =====================================================================
h("6 O QUE MELHORAR NA ENTREGA (Q2)")
para("Decompondo o prazo em duas etapas — postagem (responsabilidade do vendedor) e transporte (responsabilidade "
     "da transportadora) — o gargalo fica evidente: a postagem é praticamente constante (~2 dias) em todo o país, "
     "enquanto o transporte salta de 6 dias no Sudeste para 14 dias no Nordeste (Figura 5).")
figura("geo_q2_gargalo_entrega.png", "Decomposição do prazo: o gargalo é o transporte, não o vendedor")
para("A origem da oferta confirma o mecanismo: pedidos de vendedor local entregam ~5 dias mais rápido no agregado. "
     "No Nordeste, o pedido local leva 11 dias (91% no prazo) contra 17 dias do não-local (87%); no Centro-Oeste, "
     "7 dias (98%) contra 13 (93%) — Figura 6.")
figura("geo_q2_local_vs_naolocal.png", "Prazo e cumprimento de promessa: vendedor local vs. não-local")
para("Por que o Nordeste demora mesmo quando “local”? Porque “local” é a mesma região, e o Nordeste é "
     "geograficamente enorme. A maioria dos pedidos locais do NE cruza estados (ex.: BA→MA, ~13,5 dias). Quando "
     "vendedor e cliente estão no mesmo estado, o NE entrega em 7 dias — praticamente igual ao Sudeste (6,5 dias). "
     "Ou seja: o problema é a distância, não o vendedor (Figura 7).")
figura("geo_q2_distancia_ne.png", "Nordeste por distância: same-state entrega como o Sudeste")

# =====================================================================
# 7. SATISFAÇÃO + WELCH
# =====================================================================
h("7 SATISFAÇÃO E A PROVA ESTATÍSTICA")
para("A satisfação segue diretamente o cumprimento da promessa de entrega. A baseline regional posiciona o "
     "Nordeste como a maior alavanca: menor nota média (3,97) e maior proporção de detratores (16,7%) — Figura 8.")
figura("geo_satisfacao_promessa.png", "Satisfação × cumprimento da promessa, por faixa de atraso e por UF")
para("Para confirmar a causalidade, aplicou-se o teste de Welch (robusto a variâncias e tamanhos amostrais "
     "desiguais) comparando a nota de pedidos no prazo versus atrasados. Em todos os recortes, a queda é de "
     "aproximadamente 2 estrelas, com p-valor praticamente nulo — diferença estatisticamente significante "
     "(Tabela 1).")
tabela(["Recorte", "Nota no prazo", "Nota atrasado", "Queda", "p-valor"],
       [["Nacional", "4,29", "2,27", "−2,02", "≈ 0"],
        ["Nordeste", "4,23", "2,18", "−2,04", "1,7 × 10⁻²⁵¹"],
        ["Centro-Oeste", "4,26", "2,18", "−2,08", "9,2 × 10⁻⁸⁶"]],
       [3.5, 3.0, 3.0, 2.5, 4.0])
centro("Tabela 1 — Teste de Welch: nota no prazo vs. atrasado", 10, bold=True, before=0, after=2)
centro("Fonte: Elaborado pelos autores (2026).", 10, after=8)

# =====================================================================
# 8. ADENDO NÃO-ENTREGUES
# =====================================================================
h("8 ADENDO: OS PEDIDOS NÃO ENTREGUES")
para("Dos 99.441 pedidos da base, 2.963 (3,0%) não foram entregues e ficaram fora da análise. Não se trata de "
     "truncamento de dados — são falhas reais (idade mediana de 9 a 13 meses, prazo prometido vencido e notas de "
     "1,3 a 2,0). Distribuem-se em quatro modos de falha (Figura 9).")
figura("geo_adendo_status.png", "Pedidos entregues (usados) vs. não entregues (excluídos), por status")
bullet("Perdido no transporte (shipped, 37%): pago e postado, nunca chegou — o mesmo elo frágil do transporte; "
       "pior no Nordeste (taxa de não entrega de 3,73% vs. 2,36% no Sul).")
bullet("Sem estoque (unavailable, 21%): pago, produto inexistente — pior nota (1,5); falha de catálogo do seller.")
bullet("Travado no funil (invoiced/processing, 21%): pago e nunca despachado — falha de fulfillment.")
bullet("Cancelado (21%): majoritariamente antes da postagem; 12% já postados.")
para("O impacto é de ~R$ 370,1 mil em mercadoria e R$ 53,6 mil em frete em receita em risco, apontando um segundo "
     "problema, de natureza operacional (integridade de estoque e SLA de fulfillment), ao lado do geográfico.")

# =====================================================================
# 9. SIMULAÇÃO
# =====================================================================
h("9 SIMULAÇÃO DO CENÁRIO DE DESCENTRALIZAÇÃO (30%)")
para("Para dimensionar o retorno, modelou-se a migração de 30% dos pedidos hoje despachados de fora da região para "
     "sourcing local, aplicando aos pedidos migrados o desempenho já observado nos pedidos locais. Trata-se de um "
     "cenário-base, não de previsão de precisão (Tabela 2).")
tabela(["Região", "Pedidos migrados (30%)", "Atrasos evitados", "Economia de frete"],
       [["Nordeste", "2.584", "92", "R$ 35.812"],
        ["Centro-Oeste", "1.625", "77", "R$ 15.429"],
        ["Total", "4.209", "169", "R$ 51.241"]],
       [4.0, 4.5, 3.5, 4.0])
centro("Tabela 2 — Simulação de descentralização de 30% (NE + CO)", 10, bold=True, before=0, after=2)
centro("Fonte: Elaborado pelos autores (2026).", 10, after=8)
para("A economia de frete financia a própria atração de sellers locais; e os pedidos migrados ganham cerca de "
     "+0,3 estrela de nota média, elevando a satisfação e, por consequência, a probabilidade de recompra.")

# =====================================================================
# 10. RECOMENDAÇÕES
# =====================================================================
h("10 RECOMENDAÇÕES")
para("Foco em Nordeste e Centro-Oeste (o Norte é reconhecido como restrição geográfica — rios e carência de "
     "estradas — fora do alcance de estratégias privadas; o Sul deve ser mantido). A métrica-norte é a satisfação "
     "média regional (Tabela 3).")
tabela(["#", "Ação", "Resolve", "Meta", "Prazo"],
       [["1", "Fulfillment distribuído nos estados-âncora (BA, PE, CE) — não um hub único", "Q1 + Q2", "NE 3,97→4,15; no prazo 87%→93%", "90–180d"],
        ["2", "Recrutar sellers locais NE/CO nas categorias que sobre-indexam", "Q1", "NE 56→150+ sellers", "6–12m"],
        ["3", "Transportadora regional dedicada (corredor e intra-NE)", "Q2", "CO 93,6%→95%", "90d"],
        ["4", "Bundles regionais nas categorias de sobre-indexação", "Q1", "↑ itens/pedido", "30–60d"],
        ["5", "Recalibrar a data prometida por rota (tático)", "Q2", "↓ quebra de promessa", "30d"]],
       [0.8, 6.3, 2.0, 4.0, 2.0])
centro("Tabela 3 — Recomendações com metas de satisfação", 10, bold=True, before=0, after=2)
centro("Fonte: Elaborado pelos autores (2026).", 10, after=8)

# =====================================================================
# 11. CONSIDERAÇÕES FINAIS
# =====================================================================
h("11 CONSIDERAÇÕES FINAIS")
para("O mapa de receita da Olist não é um destino, e sim a consequência de onde a oferta está localizada. Nordeste "
     "e Centro-Oeste compram pouco e recebem com atraso pela mesma razão — a ausência de vendedores próximos. A "
     "prova está nos próprios dados: onde o vendedor é do mesmo estado, o Nordeste entrega em sete dias, igual ao "
     "Sudeste. Aproximar a oferta — de forma distribuída pelos estados de maior demanda, e não em um hub único — "
     "converte a demanda premium já existente e cumpre a promessa de entrega, medido pelo indicador que mais "
     "importa para a marca: a satisfação dessas regiões.")

# =====================================================================
# REFERÊNCIAS
# =====================================================================
page_break()
h("REFERÊNCIAS", size=13)
for ref in [
    "KAGGLE. Brazilian E-Commerce Public Dataset by Olist. Disponível em: https://www.kaggle.com/datasets/olistbr/brazilian-ecommerce. Acesso em: jun. 2026.",
    "CHOPRA, S.; MEINDL, P. Gestão da cadeia de suprimentos: estratégia, planejamento e operação. 6. ed. São Paulo: Pearson, 2016.",
    "PROVOST, F.; FAWCETT, T. Data science para negócios. Rio de Janeiro: Alta Books, 2016.",
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE; p.paragraph_format.space_after = Pt(8)
    _set_font(p.add_run(ref), 11)

add_page_numbers()
OUT.parent.mkdir(exist_ok=True, parents=True)
doc.save(str(OUT))
print(f"OK -> {OUT.relative_to(ROOT)}")
